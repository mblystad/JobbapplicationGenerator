import customtkinter as ctk
from tkinter import messagebox, filedialog
from PIL import Image, ImageTk
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from pypdf import PdfReader
import google.generativeai as genai
from datetime import datetime
import os

# Configure CustomTkinter appearance
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# Google Generative AI API key setup
API_KEY = "YOUR API KEY HERE"
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")


# Function to generate text via Google Generative AI
def generate_google_ai_text(prompt):
    response = model.generate_content(prompt)
    return response.text


# Function to scrape job descriptions from Finn.no
def scrape_job_description():
    job_url = job_url_entry.get()
    if not job_url:
        messagebox.showwarning("Manglende URL", "Vennligst skriv inn URL til jobbannonsen.")
        return None
    try:
        response = requests.get(job_url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')

        # Extract job details
        job_title_tag = soup.find("h2", class_="t2") or soup.find("h1", class_="t3") or soup.find("h1", class_="")
        job_title = job_title_tag.get_text(strip=True) if job_title_tag else "Jobbtittel ikke funnet"

        frist_tag = soup.find("li", string="Frist") or soup.find("span", string="Frist")
        frist = frist_tag.find("span").get_text(strip=True) if frist_tag else "Ikke funnet"

        ansettelse_tag = soup.find("li", string="Ansettelsesform") or soup.find("span", string="Ansettelsesform")
        ansettelsesform = ansettelse_tag.find("span").get_text(strip=True) if ansettelse_tag else "Ikke funnet"

        description_section = soup.find('section', class_='import-decoration') or soup.find('div',
                                                                                            class_='import-decoration')
        job_description = "\n".join(
            [p.get_text(strip=True) for p in description_section.find_all('p')]) if description_section else ""

        job_responsibilities = []
        responsibilities_heading = soup.find("h3", string="Oppgaver") or soup.find("h3",
                                                                                   string="Arbeidsoppgaver") or soup.find(
            "strong", string="Hva går jobben ut på?")
        if responsibilities_heading:
            ul_tag = responsibilities_heading.find_next("ul")
            if ul_tag:
                job_responsibilities = [li.get_text(strip=True) for li in ul_tag.find_all("li")]

        qualifications = []
        qualifications_heading = soup.find("h3", string="Kvalifikasjoner") or soup.find("strong",
                                                                                        string="Vi ser etter deg som har ...")
        if qualifications_heading:
            ul_tag = qualifications_heading.find_next("ul")
            if ul_tag:
                qualifications = [li.get_text(strip=True) for li in ul_tag.find_all("li")]

        job_info = {
            "title": job_title,
            "deadline": frist,
            "employment_type": ansettelsesform,
            "description": job_description,
            "responsibilities": job_responsibilities,
            "qualifications": qualifications,
            "url": job_url  # Store the URL for later reference
        }
        return job_info

    except Exception as e:
        messagebox.showerror("Feil", f"Kunne ikke skrape URL. Feil: {str(e)}")
        return None


# Function to load CV content from the file path in the entry
def load_cv_content(file_path):
    if file_path:
        if file_path.endswith('.pdf'):
            reader = PdfReader(file_path)
            return "\n".join([page.extract_text() for page in reader.pages])
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
    return None


# Select and display CV file path in entry
def select_cv_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("PDF filer", "*.pdf"), ("Word filer", "*.docx"), ("Tekstfiler", "*.txt")])
    if file_path:
        cv_path_entry.delete(0, ctk.END)
        cv_path_entry.insert(0, file_path)


# Function to upload an image file and display it
def upload_image():
    file_path = filedialog.askopenfilename(filetypes=[("Bilder", "*.jpg;*.jpeg;*.png")])
    if file_path:
        uploaded_image_path.set(file_path)

        # Load and resize the image to display in the GUI
        img = Image.open(file_path)
        img = img.resize((150, 150))  # Resize for display
        img_tk = ImageTk.PhotoImage(img)

        # Update the label with the new image
        image_label.configure(image=img_tk)
        image_label.image = img_tk  # Keep a reference to avoid garbage collection


# Generate job application using Google Gemini AI
def generate_job_application(job_info, cv_content):
    user_name = navn_entry.get()
    tone = tone_choice.get()
    tone_prompt = "Skriv en jobbsøknad til stillingen, bruk info fra cv"

    prompt = f"{tone_prompt}\nNavn: {user_name}\nJobbeskrivelse: {job_info['description']}\nMin CV: {cv_content or 'Ingen CV tilgjengelig'}"
    return generate_google_ai_text(prompt)


# Generate and save the job application
def generate_and_save_application():
    job_info = scrape_job_description()
    if not job_info:
        return

    cv_content = load_cv_content(cv_path_entry.get())
    image_path = uploaded_image_path.get()
    application_text = generate_job_application(job_info, cv_content)
    doc = Document()

    # Add image if available
    if image_path:
        doc.add_picture(image_path, width=Inches(1.5))

    # Personal info section
    today_date = datetime.now().strftime("%d.%m.%Y")
    personal_info = f"{today_date}\n\n{navn_entry.get()}\n{adresse_entry.get()}\n{tlf_entry.get()}\n{epost_entry.get()}"
    doc.add_paragraph(personal_info).alignment = 2

    # Add job application content
    doc.add_heading(f"Søknad til {job_info['title']}", level=1)
    doc.add_paragraph(application_text)

    # Page break and job info section
    doc.add_page_break()
    doc.add_heading("Jobbinformasjon", level=2)
    doc.add_paragraph(f"Tittel: {job_info['title']}")
    doc.add_paragraph(f"Frist: {job_info['deadline']}")
    doc.add_paragraph(f"Ansettelsesform: {job_info['employment_type']}")
    doc.add_heading("Jobbeskrivelse", level=2)
    doc.add_paragraph(job_info['description'])
    doc.add_heading("Ansvar", level=2)
    for responsibility in job_info['responsibilities']:
        doc.add_paragraph(f"- {responsibility}")
    doc.add_heading("Kvalifikasjoner", level=2)
    for qualification in job_info['qualifications']:
        doc.add_paragraph(f"- {qualification}")

    doc.add_paragraph("\n\nFor mer informasjon, besøk annonsen her:")
    doc.add_paragraph(job_info['url'])

    # Save document
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"Søknad - {job_info['title']}.docx",
                                             filetypes=[("Word filer", "*.docx")])
    if save_path:
        doc.save(save_path)
        messagebox.showinfo("Suksess", "Jobbsøknad generert og lagret!")


# Set up main window in CustomTkinter
root = ctk.CTk()
root.title("Jobbsøknad Generator")
root.geometry("500x800")  # Increase height to accommodate the image

# Logo
logo_image = ctk.CTkImage(dark_image=Image.open("logo.png").resize((300, 300)), size=(150, 150))
logo_label = ctk.CTkLabel(root, image=logo_image, text="")
logo_label.pack(pady=10)

# Job announcement URL entry
job_url_label = ctk.CTkLabel(root, text="Skriv inn URL til jobbannonsen:")
job_url_label.pack(pady=(10, 0))
job_url_entry = ctk.CTkEntry(root, width=300)
job_url_entry.pack(pady=5)

# CV file path entry and selection button
cv_path_label = ctk.CTkLabel(root, text="CV-fil:")
cv_path_label.pack(pady=(10, 0))
cv_path_entry = ctk.CTkEntry(root, width=300)
cv_path_entry.pack(pady=5)
cv_select_button = ctk.CTkButton(root, text="Velg CV-fil", command=select_cv_file)
cv_select_button.pack(pady=10)

# Image preview label (will show selected image)
image_label = ctk.CTkLabel(root, text="")
image_label.pack(pady=5)

# Image upload button
image_button = ctk.CTkButton(root, text="Last opp bilde", command=upload_image)
image_button.pack()


# Fields for personal information
def create_entry(parent, label_text):
    frame = ctk.CTkFrame(parent)
    frame.pack(pady=5)
    label = ctk.CTkLabel(frame, text=label_text)
    label.pack(side="left", padx=5)
    entry = ctk.CTkEntry(frame, width=200)
    entry.pack(side="right", padx=5)
    return entry


navn_entry = create_entry(root, "Navn:")
adresse_entry = create_entry(root, "Adresse:")
tlf_entry = create_entry(root, "Telefon:")
epost_entry = create_entry(root, "E-post:")

# Tone selection dropdown
tone_choice_label = ctk.CTkLabel(root, text="Velg tonetilnærming:")
tone_choice_label.pack(pady=(20, 0))
tone_choice = ctk.StringVar(value="Standard")
tone_dropdown = ctk.CTkOptionMenu(root, variable=tone_choice, values=["Standard", "Formell", "Uformell"])
tone_dropdown.pack(pady=5)

# Generate application button
generate_button = ctk.CTkButton(root, text="Generer søknad", command=generate_and_save_application)
generate_button.pack(pady=20)

# Variable to store image path
uploaded_image_path = ctk.StringVar()
root.mainloop()
